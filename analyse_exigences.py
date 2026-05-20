#!/usr/bin/env python3
"""
analyse_exigences.py
--------------------
Serveur Flask local — analyse de demandes clients en contexte naval/sous-marin.
Supporte Ollama (llama3 / mistral) et l'API Anthropic Claude (haiku / sonnet).
Exports : JSON, Word (.docx), Excel (.xlsx)

Usage : python analyse_exigences.py  →  http://localhost:5000
Clé API Claude : définir ANTHROPIC_API_KEY dans .env ou variable d'environnement.
"""

import io
import json
import re
import datetime
import ollama
from flask import Flask, request, jsonify, send_from_directory, send_file
from pathlib import Path


# pypdf pour extraction PDF (optionnel)
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

# python-docx
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# openpyxl
from openpyxl import Workbook
from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side)
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

# json-repair (optionnel)
try:
    from json_repair import repair_json
    JSON_REPAIR_AVAILABLE = True
except ImportError:
    JSON_REPAIR_AVAILABLE = False

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
DEFAULT_MODEL = "mistral"
HOST = "127.0.0.1"
PORT = 5000
OUTPUT_DIR = Path("exports")
OUTPUT_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------------------------
# Prompt système
# ---------------------------------------------------------------------------
SYSTEM_PROMPT = """Tu es une API JSON. Tu produis UNIQUEMENT du JSON valide, sans explication, sans markdown, aucun texte avant ou après l'objet JSON.

Tu es un Ingénieur Système Senior spécialisé dans les systèmes de communication des sous-marins militaires et navals :
- Communications externes immergées : VLF/ELF (stations Rosnay/Cutler/Harold Holt), SATCOM périscopique (SHF/EHF/UHF), liaisons acoustiques sous-marines longue portée
- Communications internes : réseaux de combat embarqués, bus militaires (MIL-STD-1553B, STANAG 3910), intercom, systèmes de commandement
- Standards : STANAG 4204/4206, MIL-STD-188-110C, MIL-STD-461G, MIL-STD-810H, DEF STAN 02-713, SDIP-27, IEC 60945
- Contraintes sous-marines : étanchéité (IP68, pression hydrostatique), EMI/EMC, chocs MIL-S-901D, corrosion marine

Tu appliques la norme ISO 29148:2018 pour la rédaction des exigences système.

APPROCHE D'ANALYSE EN 3 NIVEAUX — FONDAMENTALE
Les demandes client sont souvent incomplètes. Ne te limite jamais à extraire littéralement le texte.

NIVEAU 1 — Explicite (origine: "Explicite")
  UNIQUEMENT si le texte source contient MOT POUR MOT l'information (valeur, standard, action précise).
  Tu DOIS pouvoir copier une phrase réelle du texte en justification.
  justification = extrait textuel réel entre guillemets «...»

  EXEMPLES CORRECTS :
    Texte: "débit minimal de 500 bps sur 5 km"
    → texte EF: "Le système DOIT assurer un débit minimal de 500 bps sur une portée de 5 km."
    → origine: "Explicite", justification: "«débit minimal de 500 bps sur 5 km»"

  RÈGLE CRITIQUE : Si le texte est vague ("définir la communication interne", "besoin d'un système..."),
  il N'Y A PAS d'Explicite possible pour des valeurs numériques. Ne force pas un "Explicite" sur un texte vague.

NIVEAU 2 — Déduite (origine: "Déduite")
  Logiquement nécessaire à partir d'une contrainte explicite. Raisonnement technique rigoureux.
  N'invente AUCUNE valeur — déduis depuis une loi physique ou règle technique connue.
  justification = raisonnement (ex: "Profondeur 400 m mentionnée → 1 bar ≈ 10 m → 40 bar minimum")

  EXEMPLES CORRECTS :
    Texte "profondeur 400 m" → Déduite: "résistance à 40 bar" / justification: "400 m → 40 bar (loi de Pascal)"
    Texte "sous-marin" seul → Déduite: "étanchéité IP68" / justification: "Environnement immergé → étanchéité IP68 minimale"
    Texte "eau de mer" → Déduite: "résistance corrosion NaCl 3,5%" / justification: "Milieu marin → corrosion saline"
    Texte vague "communication interne" → Déduite: "redondance réseau" / justification: "Système critique embarqué → redondance standard"

NIVEAU 3 — Hypothèse (origine: "Hypothèse")
  Standard du domaine naval/militaire, non mentionné dans le texte mais attendu.
  justification = "Hypothèse domaine : [raison]. A valider avec le client."

  EXEMPLES CORRECTS :
    Texte vague → Hypothèse: "Conformité MIL-STD-461G" / justification: "Hypothèse domaine : système militaire embarqué → CEM militaire. A valider."
    Texte vague → Hypothèse: "Interface MIL-STD-1553B" / justification: "Hypothèse domaine : sous-marin OTAN → bus militaire standard. A valider."

RÈGLE DE MIX pour texte VAGUE (peu de détails) :
  → 0 ou 1 Explicite MAX, majorité Déduites + Hypothèses
  → Ne génère JAMAIS "Explicite" si tu n'as pas de citation réelle à mettre en justification

RÈGLE DE MIX pour texte DÉTAILLÉ (valeurs numériques présentes) :
  → Génère des Explicites en proportion du détail fourni, complétées par Déduites + Hypothèses

RÈGLE ANTI-INVENTION : Pour Déduite et Hypothèse, si la valeur exacte n'est pas déductible,
écris "à définir lors de la revue des exigences" — jamais une valeur inventée.

STRUCTURE JSON EXACTE :
{
  "perimetre": {
    "systeme_sous_analyse": "nom court du système",
    "acteurs_externes": ["Opérateur Radio", "Autorité de commandement OTAN"],
    "milieux_externes": ["Milieu marin immergé", "Eau de mer 3,5% NaCl"]
  },
  "besoin_fondamental": "phrase unique décrivant la mission de communication avec contexte opérationnel",
  "exigences_fonctionnelles": [
    {
      "texte": "Le système DOIT recevoir les messages OTAN en bande VLF (3-30 kHz) a une profondeur minimale de X metres.",
      "origine": "Explicite",
      "justification": "«citation exacte du texte source»",
      "statut": "A confirmer"
    },
    {
      "texte": "Le systeme DOIT resister a une pression hydrostatique minimale de 40 bar.",
      "origine": "Déduite",
      "justification": "Profondeur operationnelle de 400 m mentionnee → 1 bar ≈ 10 m d'eau → pression min = 40 bar.",
      "statut": "A confirmer"
    },
    {
      "texte": "Le systeme DOIT etre conforme au standard MIL-STD-461G pour la compatibilite electromagnetique.",
      "origine": "Hypothèse",
      "justification": "Hypothese domaine : systeme militaire embarque sur sous-marin → conformite CEM militaire standard. A valider avec le client.",
      "statut": "A confirmer"
    }
  ],
  "exigences_non_fonctionnelles": [
    {
      "type": "Performance",
      "texte": "Le systeme DOIT atteindre une sensibilite minimale de X dBm.",
      "origine": "Explicite",
      "justification": "«citation exacte»",
      "statut": "A confirmer"
    },
    {
      "type": "Hypothèse",
      "texte": "Le systeme DOIT presenter une MTBF minimale a definir lors de la revue des exigences.",
      "origine": "Hypothèse",
      "justification": "Hypothese domaine : systeme militaire critique → exigence de fiabilite standard naval. A valider avec le client.",
      "statut": "A confirmer"
    }
  ],
  "contraintes_techniques": {
    "acoustique": "valeurs acoustiques ou Non specifie",
    "frequences": "bandes de frequences (ex: VLF 3-30 kHz) ou Non specifie",
    "etancheite_pression": "indice IP et pression en bar ou Non specifie",
    "debit_latence": "debit en bps/kbps et latence en ms ou Non specifie",
    "environnement_physique": "temperature C, salinite %, vibrations g ou Non specifie",
    "autres": ["normes applicables si mentionnees"]
  },
  "mots_cles_domaine": ["motcle1", "motcle2", "motcle3", "motcle4", "motcle5"]
}

RÈGLES ABSOLUES :
1. origine DOIT être exactement : "Explicite" ou "Déduite" ou "Hypothèse"
2. statut DOIT être : "A confirmer"
3. Chaque texte DOIT commencer par "Le système DOIT" ou "Le systeme DOIT"
4. type ENF : Performance | Fiabilite | Interface | Physique | Securite
5. Génère entre 5 et 9 exigences fonctionnelles (mix obligatoire des 3 niveaux)
6. Génère entre 2 et 4 exigences non-fonctionnelles
7. Toutes les valeurs en FRANÇAIS
8. Si document de contexte fourni : valeurs numériques DOIVENT prioritairement en provenir
9. PRODUIS UNIQUEMENT l'objet JSON, commençant par { et se terminant par }
"""

# ---------------------------------------------------------------------------
# Parsing JSON robuste
# ---------------------------------------------------------------------------
def parse_json_robust(raw_text: str) -> dict:
    text = re.sub(r"```json\s*", "", raw_text)
    text = re.sub(r"```\s*", "", text).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    first, last = text.find("{"), text.rfind("}")
    if first != -1 and last != -1 and last > first:
        candidate = text[first:last + 1]
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            pass
        if JSON_REPAIR_AVAILABLE:
            try:
                return json.loads(repair_json(candidate))
            except Exception:
                pass
    if JSON_REPAIR_AVAILABLE:
        try:
            return json.loads(repair_json(text))
        except Exception:
            pass
    raise json.JSONDecodeError("Impossible de parser le JSON", text, 0)

# ---------------------------------------------------------------------------
# Flask app
# ---------------------------------------------------------------------------
app = Flask(__name__, static_folder=".")


@app.route("/")
def index():
    return send_from_directory(".", "index.html")


@app.route("/analyse", methods=["POST"])
def analyse():
    data    = request.get_json(force=True)
    demande = data.get("demande", "").strip()
    model   = data.get("model", DEFAULT_MODEL).strip()
    context = data.get("context", "").strip()
    if not demande:
        return jsonify({"error": "Demande vide."}), 400

    parts = []
    if context:
        parts.append(f"DOCUMENT DE CONTEXTE (référence technique — valeurs numériques prioritaires) :\n{context[:20000]}")
    parts.append(f"DEMANDE CLIENT À ANALYSER :\n\n{demande}\n\nProduis l'objet JSON maintenant :")
    user_message = "\n\n---\n\n".join(parts)

    try:
        response = ollama.chat(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user",   "content": user_message},
            ],
            options={"temperature": 0.1, "top_p": 0.9,
                     "num_predict": 4096, "num_ctx": 16384}
        )
        raw = (response.message.content
               if hasattr(response, "message")
               else response["message"]["content"])
        result = parse_json_robust(raw)
        result["_meta"] = {
            "model":          model,
            "timestamp":      datetime.datetime.now().isoformat(),
            "demande_source": demande[:200] + ("..." if len(demande) > 200 else ""),
            "demande_full":   demande,
            "context_file":   data.get("context_filename", ""),
        }
        return jsonify(result)
    except json.JSONDecodeError as e:
        return jsonify({"error": f"JSON invalide : {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"Erreur inattendue : {str(e)}"}), 500


@app.route("/extract-context", methods=["POST"])
def extract_context():
    """Extrait le texte d'un PDF ou fichier texte pour l'utiliser comme contexte."""
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "Aucun fichier reçu."}), 400

    fname = (file.filename or "").lower()

    if fname.endswith(".pdf"):
        if not PYPDF_AVAILABLE:
            return jsonify({"error": "Package 'pypdf' non installé. Lancez : pip install pypdf"}), 500
        try:
            reader = pypdf.PdfReader(file)
            text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        except Exception as e:
            return jsonify({"error": f"Impossible de lire le PDF : {str(e)}"}), 500
    elif fname.endswith((".txt", ".md")):
        text = file.read().decode("utf-8", errors="ignore").strip()
    else:
        return jsonify({"error": "Format non supporté. Utilisez PDF, TXT ou MD."}), 400

    MAX_CHARS = 80000
    truncated = len(text) > MAX_CHARS
    text = text[:MAX_CHARS]
    return jsonify({"text": text, "chars": len(text), "truncated": truncated})


@app.route("/affiner", methods=["POST"])
def affiner():
    """Affine les exigences existantes pour les rendre plus précises (ISO 29148)."""
    data   = request.get_json(force=True)
    result = data.get("result")
    model  = data.get("model", DEFAULT_MODEL).strip()
    if not result:
        return jsonify({"error": "Aucun résultat à affiner."}), 400

    refine_msg = (
        "Voici les exigences système extraites. Améliore-les pour qu'elles soient "
        "plus précises, mesurables et conformes à la norme ISO 29148 "
        "(atomiques, vérifiables, non-ambiguës, avec valeurs chiffrées quand possible). "
        "Conserve la structure JSON exacte et le sens de chaque exigence. "
        "Améliore uniquement la qualité rédactionnelle.\n\n"
        f"Objet JSON actuel :\n{json.dumps(result, ensure_ascii=False, indent=2)}\n\n"
        "Retourne l'objet JSON amélioré :"
    )

    try:
        response = ollama.chat(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user",   "content": refine_msg},
            ],
            options={"temperature": 0.2, "top_p": 0.9,
                     "num_predict": 4096, "num_ctx": 16384}
        )
        raw = (response.message.content
               if hasattr(response, "message")
               else response["message"]["content"])
        refined = parse_json_robust(raw)
        refined["_meta"] = result.get("_meta", {})
        refined["_meta"]["refined_at"] = datetime.datetime.now().isoformat()
        return jsonify(refined)
    except json.JSONDecodeError as e:
        return jsonify({"error": f"JSON invalide : {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"Erreur inattendue : {str(e)}"}), 500


@app.route("/export", methods=["POST"])
def export_json():
    """Export JSON brut sur le disque local."""
    data   = request.get_json(force=True)
    result = data.get("result")
    if not result:
        return jsonify({"error": "Rien à exporter."}), 400
    ts       = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"exigences_{ts}.json"
    with open(OUTPUT_DIR / filename, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    return jsonify({"message": "Fichier exporté.", "filename": filename})


# ===========================================================================
# EXPORT WORD (.docx)
# ===========================================================================
@app.route("/export/word", methods=["POST"])
def export_word():
    """Génère et renvoie un fichier .docx professionnel en téléchargement."""
    data   = request.get_json(force=True)
    result = data.get("result")
    if not result:
        return jsonify({"error": "Rien à exporter."}), 400

    p   = result.get("perimetre", {})
    ef  = result.get("exigences_fonctionnelles", [])
    ct  = result.get("contraintes_techniques", {})
    kw  = result.get("mots_cles_domaine", [])
    meta = result.get("_meta", {})

    doc = DocxDocument()

    # ---- Marges de page ----
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ---- Couleurs thème ----
    NAVY    = RGBColor(0x05, 0x37, 0x61)
    CYAN    = RGBColor(0x00, 0x9B, 0xC8)
    ORANGE  = RGBColor(0xC0, 0x50, 0x10)
    GRAY_BG = "D9E8F3"
    GRAY_HD = "053761"
    WHITE   = "FFFFFF"

    def set_cell_bg(cell, hex_color):
        """Applique une couleur de fond à une cellule."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def cell_border(cell, border_color="009BC8", size=6):
        """Ajoute des bordures à une cellule."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side_name in ['top','left','bottom','right']:
            side = OxmlElement(f'w:{side_name}')
            side.set(qn('w:val'),   'single')
            side.set(qn('w:sz'),    str(size))
            side.set(qn('w:space'), '0')
            side.set(qn('w:color'), border_color)
            tcBorders.append(side)
        tcPr.append(tcBorders)

    def add_heading(text, level=1, color=NAVY):
        p_obj = doc.add_paragraph()
        p_obj.paragraph_format.space_before = Pt(14 if level == 1 else 8)
        p_obj.paragraph_format.space_after  = Pt(4)
        run = p_obj.add_run(text)
        run.bold      = True
        run.font.size = Pt(16 if level == 1 else 12)
        run.font.color.rgb = color
        if level == 1:
            # Ligne de soulignement via border bottom sur le paragraphe
            pPr = p_obj._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'),   'single')
            bottom.set(qn('w:sz'),    '6')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), '009BC8')
            pBdr.append(bottom)
            pPr.append(pBdr)
        return p_obj

    # ===== TITRE PRINCIPAL =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after  = Pt(4)
    r = title_para.add_run("FICHE D'EXIGENCES SYSTÈME")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(2)
    r2 = subtitle_para.add_run(p.get("systeme_sous_analyse", "Système non spécifié").upper())
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = CYAN

    # Métadonnées
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_after = Pt(12)
    ts_str = ""
    if meta.get("timestamp"):
        try:
            ts_str = datetime.datetime.fromisoformat(meta["timestamp"]).strftime("%d/%m/%Y %H:%M")
        except Exception:
            ts_str = meta["timestamp"]
    r3 = meta_para.add_run(f"Généré le {ts_str}  |  Modèle : {meta.get('model','?')}  |  SYREQ v1.0")
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0x88, 0xAA, 0xBB)
    r3.italic = True

    doc.add_paragraph()  # espace

    # ===== 0. DEMANDE CLIENT =====
    demande_full = meta.get("demande_full") or meta.get("demande_source") or ""
    context_file = meta.get("context_file", "")
    if demande_full:
        add_heading("Demande Client Analysée", level=2, color=CYAN)
        dem_para = doc.add_paragraph()
        dem_para.paragraph_format.left_indent  = Cm(0.5)
        dem_para.paragraph_format.space_after  = Pt(4)
        dem_para.paragraph_format.space_before = Pt(2)
        r_dem = dem_para.add_run(demande_full)
        r_dem.font.size = Pt(10)
        r_dem.italic = True
        if context_file:
            ctx_para = doc.add_paragraph()
            ctx_para.paragraph_format.space_after = Pt(8)
            r_ctx = ctx_para.add_run(f"Document de contexte : {context_file}")
            r_ctx.font.size = Pt(8)
            r_ctx.font.color.rgb = RGBColor(0x88, 0xAA, 0xBB)
            r_ctx.italic = True
        doc.add_paragraph()

    # ===== 1. PÉRIMÈTRE =====
    add_heading("1. PÉRIMÈTRE DU SYSTÈME")

    # Tableau acteurs / milieux
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr_cells = tbl.rows[0].cells
    for i, txt in enumerate(["ACTEURS EXTERNES", "MILIEUX EXTERNES"]):
        hdr_cells[i].text = txt
        set_cell_bg(hdr_cells[i], GRAY_HD)
        cell_border(hdr_cells[i])
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        run.font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contenu acteurs
    acteurs = p.get("acteurs_externes", [])
    milieux = p.get("milieux_externes", [])
    max_rows = max(len(acteurs), len(milieux), 1)

    # Ajouter les lignes nécessaires (on en a déjà 2 après le header)
    while len(tbl.rows) - 1 < max_rows:
        tbl.add_row()

    for idx in range(max_rows):
        row = tbl.rows[idx + 1]
        act_text = acteurs[idx] if idx < len(acteurs) else ""
        mil_text = milieux[idx] if idx < len(milieux) else ""
        for col_i, txt in enumerate([act_text, mil_text]):
            cell = row.cells[col_i]
            cell.text = txt
            cell.paragraphs[0].runs[0].font.size = Pt(10) if txt else Pt(10)
            set_cell_bg(cell, GRAY_BG if idx % 2 == 0 else WHITE)
            cell_border(cell, "AACCDD", 4)

    # Supprimer la ligne vide initiale (row index 2 si max_rows=1)
    if len(tbl.rows) > max_rows + 1:
        tbl._tbl.remove(tbl.rows[-1]._tr)

    doc.add_paragraph()

    # ===== 2. BESOIN FONDAMENTAL =====
    add_heading("2. BESOIN FONDAMENTAL")
    besoin_para = doc.add_paragraph()
    besoin_para.paragraph_format.left_indent  = Cm(0.8)
    besoin_para.paragraph_format.space_after  = Pt(8)
    besoin_para.paragraph_format.space_before = Pt(4)
    # Barre latérale via border left
    pPr2 = besoin_para._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'),   'single')
    left_b.set(qn('w:sz'),    '12')
    left_b.set(qn('w:space'), '8')
    left_b.set(qn('w:color'), '009BC8')
    pBdr2.append(left_b)
    pPr2.append(pBdr2)
    r_b = besoin_para.add_run(result.get("besoin_fondamental", "Non spécifié"))
    r_b.font.size = Pt(11)
    r_b.italic = True

    # helpers pour lire EF/NFR (objets ou chaînes)
    def ef_text(e):         return e.get("text", e) if isinstance(e, dict) else e
    def ef_priority(e):     return e.get("priority", "Must") if isinstance(e, dict) else "Must"
    def ef_verif(e):        return e.get("verification", "Test") if isinstance(e, dict) else "Test"

    PRIORITY_COLORS = {
        "Must":    RGBColor(0xFF, 0x3D, 0x5A),
        "Should":  RGBColor(0xFF, 0xCC, 0x00),
        "Could":   RGBColor(0x00, 0x9B, 0xC8),
        "Won't":   RGBColor(0x66, 0x66, 0x66),
    }

    def add_ef_row(tbl, idx, item_id, text_val, priority, verif):
        row = tbl.add_row()
        bg  = GRAY_BG if idx % 2 == 0 else WHITE

        id_cell = row.cells[0]
        id_cell.text = item_id
        id_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        id_cell.paragraphs[0].runs[0].bold = True
        id_cell.paragraphs[0].runs[0].font.color.rgb = CYAN
        id_cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(id_cell, bg); cell_border(id_cell, "AACCDD", 4)

        exig_cell = row.cells[1]
        exig_para = exig_cell.paragraphs[0]
        exig_para.paragraph_format.space_before = Pt(2)
        exig_para.paragraph_format.space_after  = Pt(2)
        doit_match = re.match(r"^(Le syst[eè]me DOIT\s*)", text_val, re.IGNORECASE)
        if doit_match:
            r_must = exig_para.add_run(doit_match.group(1))
            r_must.bold = True; r_must.font.color.rgb = ORANGE; r_must.font.size = Pt(10)
            r_rest = exig_para.add_run(text_val[len(doit_match.group(1)):])
            r_rest.font.size = Pt(10)
        else:
            r_all = exig_para.add_run(text_val); r_all.font.size = Pt(10)
        set_cell_bg(exig_cell, bg); cell_border(exig_cell, "AACCDD", 4)

        prio_cell = row.cells[2]
        prio_cell.text = priority
        prio_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        prio_cell.paragraphs[0].runs[0].bold = True
        prio_cell.paragraphs[0].runs[0].font.size = Pt(9)
        prio_cell.paragraphs[0].runs[0].font.color.rgb = PRIORITY_COLORS.get(priority, RGBColor(0x66,0x66,0x66))
        set_cell_bg(prio_cell, bg); cell_border(prio_cell, "AACCDD", 4)

        verif_cell = row.cells[3]
        verif_cell.text = verif
        verif_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        verif_cell.paragraphs[0].runs[0].font.size = Pt(9)
        verif_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x9B, 0xC8)
        set_cell_bg(verif_cell, bg); cell_border(verif_cell, "AACCDD", 4)

        stat_cell = row.cells[4]
        stat_cell.text = "À valider"
        stat_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        stat_cell.paragraphs[0].runs[0].font.size = Pt(9)
        stat_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        set_cell_bg(stat_cell, bg); cell_border(stat_cell, "AACCDD", 4)

    # ===== 3. EXIGENCES FONCTIONNELLES =====
    add_heading(f"3. EXIGENCES FONCTIONNELLES  ({len(ef)})")

    tbl_ef = doc.add_table(rows=1, cols=5)
    tbl_ef.style = 'Table Grid'
    tbl_ef.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_ws_ef = [600, 4800, 1000, 1200, 1000]
    for i, cell in enumerate(tbl_ef.rows[0].cells):
        cell.width = col_ws_ef[i]

    for i, txt in enumerate(["ID", "EXIGENCE", "PRIORITÉ", "VÉRIF.", "STATUT"]):
        cell = tbl_ef.rows[0].cells[i]
        cell.text = txt
        set_cell_bg(cell, GRAY_HD); cell_border(cell)
        run = cell.paragraphs[0].runs[0]
        run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, exig in enumerate(ef):
        add_ef_row(tbl_ef, idx,
                   f"EF-{str(idx+1).zfill(2)}",
                   ef_text(exig), ef_priority(exig), ef_verif(exig))

    doc.add_paragraph()

    # ===== 4. EXIGENCES NON-FONCTIONNELLES =====
    nfr = result.get("exigences_non_fonctionnelles", [])
    if nfr:
        add_heading(f"4. EXIGENCES NON-FONCTIONNELLES  ({len(nfr)})")

        tbl_nfr = doc.add_table(rows=1, cols=5)
        tbl_nfr.style = 'Table Grid'
        tbl_nfr.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, cell in enumerate(tbl_nfr.rows[0].cells):
            cell.width = col_ws_ef[i]

        for i, txt in enumerate(["ID", "EXIGENCE", "TYPE", "PRIORITÉ", "VÉRIF."]):
            cell = tbl_nfr.rows[0].cells[i]
            cell.text = txt
            set_cell_bg(cell, GRAY_HD); cell_border(cell)
            run = cell.paragraphs[0].runs[0]
            run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, item in enumerate(nfr):
            nfr_text = item.get("text", item.get("texte", "")) if isinstance(item, dict) else str(item)
            nfr_type = item.get("type", "Performance") if isinstance(item, dict) else "Performance"
            nfr_prio = item.get("priority", "Must") if isinstance(item, dict) else "Must"
            nfr_verif = item.get("verification", "Test") if isinstance(item, dict) else "Test"
            row = tbl_nfr.add_row()
            bg  = GRAY_BG if idx % 2 == 0 else WHITE

            r0 = row.cells[0]
            r0.text = f"ENF-{str(idx+1).zfill(2)}"
            r0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0.paragraphs[0].runs[0].bold = True
            r0.paragraphs[0].runs[0].font.color.rgb = CYAN
            r0.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(r0, bg); cell_border(r0, "AACCDD", 4)

            r1 = row.cells[1]
            r1.text = nfr_text; r1.paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(r1, bg); cell_border(r1, "AACCDD", 4)

            r2 = row.cells[2]
            r2.text = nfr_type
            r2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2.paragraphs[0].runs[0].font.size = Pt(9)
            r2.paragraphs[0].runs[0].font.color.rgb = ORANGE
            set_cell_bg(r2, bg); cell_border(r2, "AACCDD", 4)

            r3 = row.cells[3]
            r3.text = nfr_prio
            r3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r3.paragraphs[0].runs[0].bold = True
            r3.paragraphs[0].runs[0].font.size = Pt(9)
            r3.paragraphs[0].runs[0].font.color.rgb = PRIORITY_COLORS.get(nfr_prio, RGBColor(0x66,0x66,0x66))
            set_cell_bg(r3, bg); cell_border(r3, "AACCDD", 4)

            r4 = row.cells[4]
            r4.text = nfr_verif
            r4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r4.paragraphs[0].runs[0].font.size = Pt(9)
            r4.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x9B, 0xC8)
            set_cell_bg(r4, bg); cell_border(r4, "AACCDD", 4)

        doc.add_paragraph()

    # ===== 5. CONTRAINTES TECHNIQUES =====
    add_heading("5. CONTRAINTES TECHNIQUES")

    CT_LABELS = [
        ("acoustique",         "Acoustique"),
        ("frequences",         "Fréquences"),
        ("etancheite_pression","Étanchéité / Pression"),
        ("debit_latence",      "Débit / Latence"),
        ("environnement_physique", "Environnement physique"),
    ]

    tbl_ct = doc.add_table(rows=1, cols=2)
    tbl_ct.style = 'Table Grid'
    tbl_ct.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header CT
    for i, txt in enumerate(["CONTRAINTE", "SPÉCIFICATION"]):
        cell = tbl_ct.rows[0].cells[i]
        cell.text = txt
        set_cell_bg(cell, GRAY_HD)
        cell_border(cell)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (key, label) in enumerate(CT_LABELS):
        val = ct.get(key, "Non spécifié")
        row = tbl_ct.add_row()
        bg  = GRAY_BG if idx % 2 == 0 else WHITE

        lbl_cell = row.cells[0]
        lbl_cell.text = label
        lbl_cell.paragraphs[0].runs[0].bold = True
        lbl_cell.paragraphs[0].runs[0].font.size = Pt(10)
        lbl_cell.paragraphs[0].runs[0].font.color.rgb = NAVY
        set_cell_bg(lbl_cell, bg)
        cell_border(lbl_cell, "AACCDD", 4)

        val_cell = row.cells[1]
        val_cell.text = val
        is_na = val in ("N/A", "Non spécifié", "Non specifie")
        val_cell.paragraphs[0].runs[0].font.size = Pt(10)
        if is_na:
            val_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            val_cell.paragraphs[0].runs[0].italic = True
        set_cell_bg(val_cell, bg)
        cell_border(val_cell, "AACCDD", 4)

    # Autres contraintes
    autres = ct.get("autres", [])
    if autres:
        row = tbl_ct.add_row()
        bg  = GRAY_BG if len(CT_LABELS) % 2 == 0 else WHITE
        lbl_cell = row.cells[0]
        lbl_cell.text = "Autres"
        lbl_cell.paragraphs[0].runs[0].bold = True
        lbl_cell.paragraphs[0].runs[0].font.size = Pt(10)
        lbl_cell.paragraphs[0].runs[0].font.color.rgb = NAVY
        set_cell_bg(lbl_cell, bg)
        cell_border(lbl_cell, "AACCDD", 4)
        val_cell = row.cells[1]
        val_cell.text = " / ".join(autres)
        val_cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(val_cell, bg)
        cell_border(val_cell, "AACCDD", 4)

    doc.add_paragraph()

    # ===== 6. MOTS-CLÉS =====
    if kw:
        add_heading("6. MOTS-CLÉS DOMAINE")
        kw_para = doc.add_paragraph()
        kw_para.paragraph_format.space_after = Pt(8)
        for i, k in enumerate(kw):
            r_kw = kw_para.add_run(f" {k} ")
            r_kw.bold = True
            r_kw.font.size = Pt(9)
            r_kw.font.color.rgb = ORANGE
            if i < len(kw) - 1:
                sep = kw_para.add_run("  ·  ")
                sep.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                sep.font.size = Pt(9)

    # ===== PIED DE PAGE =====
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = footer_para.add_run("Document généré par SYREQ — Usage interne confidentiel")
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = RGBColor(0x88, 0xAA, 0xBB)
    r_f.italic = True

    # ===== Sérialisation en mémoire =====
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"exigences_{ts}.docx"

    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ===========================================================================
# EXPORT EXCEL (.xlsx)
# ===========================================================================
@app.route("/export/excel", methods=["POST"])
def export_excel():
    """Génère et renvoie un fichier .xlsx structuré en téléchargement."""
    data   = request.get_json(force=True)
    result = data.get("result")
    if not result:
        return jsonify({"error": "Rien à exporter."}), 400

    p    = result.get("perimetre", {})
    ef   = result.get("exigences_fonctionnelles", [])
    ct   = result.get("contraintes_techniques", {})
    kw   = result.get("mots_cles_domaine", [])
    meta = result.get("_meta", {})

    wb = Workbook()
    wb.remove(wb.active)  # On crée nos feuilles proprement

    # ---- Styles communs ----
    NAVY_HEX   = "053761"
    CYAN_HEX   = "009BC8"
    GREEN_HEX  = "007A4B"
    ORANGE_HEX = "C05010"
    LIGHT_BG   = "D9E8F3"
    ALT_BG     = "EEF6FB"
    WHITE_HEX  = "FFFFFF"

    def header_fill(hex_color=NAVY_HEX):
        return PatternFill("solid", fgColor=hex_color)

    def light_fill(hex_color=LIGHT_BG):
        return PatternFill("solid", fgColor=hex_color)

    def alt_fill():
        return PatternFill("solid", fgColor=ALT_BG)

    def thin_border(color="AACCDD"):
        s = Side(style="thin", color=color)
        return Border(left=s, right=s, top=s, bottom=s)

    def hdr_font(size=11, color=WHITE_HEX):
        return Font(name="Arial", bold=True, size=size, color=color)

    def body_font(size=10, bold=False, color="000000", italic=False):
        return Font(name="Arial", size=size, bold=bold, color=color, italic=italic)

    def center_align(wrap=False):
        return Alignment(horizontal="center", vertical="center", wrap_text=wrap)

    def left_align(wrap=False):
        return Alignment(horizontal="left", vertical="center", wrap_text=wrap)

    def apply_header_row(ws, row_num, headers, col_widths=None):
        for col_i, txt in enumerate(headers, start=1):
            cell = ws.cell(row=row_num, column=col_i, value=txt)
            cell.fill   = header_fill()
            cell.font   = hdr_font()
            cell.border = thin_border("009BC8")
            cell.alignment = center_align()
        if col_widths:
            for col_i, w in enumerate(col_widths, start=1):
                ws.column_dimensions[get_column_letter(col_i)].width = w

    # =========================================================
    # FEUILLE 1 : SYNTHÈSE
    # =========================================================
    ws1 = wb.create_sheet("Synthèse")
    ws1.sheet_view.showGridLines = False

    # Titre
    ws1.merge_cells("A1:D1")
    title_cell = ws1["A1"]
    title_cell.value     = "FICHE D'EXIGENCES SYSTÈME — SYREQ"
    title_cell.font      = Font(name="Arial", bold=True, size=16, color=NAVY_HEX)
    title_cell.fill      = PatternFill("solid", fgColor=LIGHT_BG)
    title_cell.alignment = center_align()
    ws1.row_dimensions[1].height = 30

    # Sous-titre système
    ws1.merge_cells("A2:D2")
    syst_cell = ws1["A2"]
    syst_cell.value     = p.get("systeme_sous_analyse", "Non spécifié").upper()
    syst_cell.font      = Font(name="Arial", bold=True, size=12, color=CYAN_HEX)
    syst_cell.fill      = PatternFill("solid", fgColor=LIGHT_BG)
    syst_cell.alignment = center_align()
    ws1.row_dimensions[2].height = 20

    # Métadonnées
    ws1.merge_cells("A3:D3")
    ts_str = ""
    if meta.get("timestamp"):
        try:
            ts_str = datetime.datetime.fromisoformat(meta["timestamp"]).strftime("%d/%m/%Y %H:%M")
        except Exception:
            ts_str = meta.get("timestamp","")
    meta_cell = ws1["A3"]
    meta_cell.value     = f"Généré le {ts_str}  |  Modèle : {meta.get('model','?')}  |  SYREQ v1.0"
    meta_cell.font      = Font(name="Arial", size=9, color="889999", italic=True)
    meta_cell.fill      = PatternFill("solid", fgColor=LIGHT_BG)
    meta_cell.alignment = center_align()
    ws1.row_dimensions[3].height = 16

    # Espace
    ws1.row_dimensions[4].height = 8

    # Bloc infos
    demande_full_xl = meta.get("demande_full") or meta.get("demande_source") or ""
    context_file_xl = meta.get("context_file", "")
    infos = [
        ("Système sous analyse", p.get("systeme_sous_analyse","Non spécifié")),
        ("Besoin fondamental",   result.get("besoin_fondamental","Non spécifié")),
        ("Acteurs externes",     ", ".join(p.get("acteurs_externes",[]))),
        ("Milieux externes",     ", ".join(p.get("milieux_externes",[]))),
        ("Mots-clés domaine",    ", ".join(kw)),
        ("Demande client",       demande_full_xl),
    ]
    if context_file_xl:
        infos.append(("Document de contexte", context_file_xl))
    for i, (label, val) in enumerate(infos, start=5):
        ws1.cell(row=i, column=1, value=label).font  = body_font(bold=True, color=NAVY_HEX)
        ws1.cell(row=i, column=1).fill               = light_fill()
        ws1.cell(row=i, column=1).border             = thin_border()
        ws1.cell(row=i, column=1).alignment          = left_align()
        ws1.merge_cells(start_row=i, start_column=2, end_row=i, end_column=4)
        val_cell = ws1.cell(row=i, column=2, value=val)
        val_cell.font      = body_font(size=10)
        val_cell.fill      = alt_fill() if i % 2 == 0 else light_fill(WHITE_HEX)
        val_cell.border    = thin_border()
        val_cell.alignment = left_align(wrap=True)
        row_h = 60 if label == "Demande client" and len(val) > 100 else 22
        ws1.row_dimensions[i].height = row_h

    ws1.column_dimensions["A"].width = 28
    ws1.column_dimensions["B"].width = 40
    ws1.column_dimensions["C"].width = 20
    ws1.column_dimensions["D"].width = 20

    # helpers EF (objets ou chaînes)
    def xl_ef_text(e):  return e.get("text", e) if isinstance(e, dict) else e
    def xl_ef_prio(e):  return e.get("priority", "Must") if isinstance(e, dict) else "Must"
    def xl_ef_verif(e): return e.get("verification", "Test") if isinstance(e, dict) else "Test"

    PRIO_COLORS_XL = {
        "Must":   "FF3D5A",
        "Should": "FFCC00",
        "Could":  "009BC8",
        "Won't":  "888888",
    }

    # =========================================================
    # FEUILLE 2 : EXIGENCES FONCTIONNELLES
    # =========================================================
    ws2 = wb.create_sheet("Exigences Fonctionnelles")
    ws2.sheet_view.showGridLines = False

    ws2.merge_cells("A1:E1")
    ef_title = ws2["A1"]
    ef_title.value     = f"EXIGENCES FONCTIONNELLES — {p.get('systeme_sous_analyse','')}"
    ef_title.font      = Font(name="Arial", bold=True, size=13, color=NAVY_HEX)
    ef_title.fill      = light_fill(LIGHT_BG)
    ef_title.alignment = center_align()
    ws2.row_dimensions[1].height = 26

    apply_header_row(ws2, 2,
        ["ID", "EXIGENCE", "PRIORITÉ", "VÉRIF.", "STATUT"],
        col_widths=[8, 58, 12, 16, 14])

    for idx, exig in enumerate(ef):
        row_num = idx + 3
        bg_fill = light_fill(LIGHT_BG) if idx % 2 == 0 else alt_fill()
        prio    = xl_ef_prio(exig)

        id_cell = ws2.cell(row=row_num, column=1, value=f"EF-{str(idx+1).zfill(2)}")
        id_cell.font      = body_font(bold=True, color=CYAN_HEX)
        id_cell.fill      = bg_fill
        id_cell.border    = thin_border()
        id_cell.alignment = center_align()

        exig_cell = ws2.cell(row=row_num, column=2, value=xl_ef_text(exig))
        exig_cell.font      = body_font(size=10)
        exig_cell.fill      = bg_fill
        exig_cell.border    = thin_border()
        exig_cell.alignment = left_align(wrap=True)

        prio_cell = ws2.cell(row=row_num, column=3, value=prio)
        prio_cell.font      = body_font(size=9, color=PRIO_COLORS_XL.get(prio, ORANGE_HEX), bold=True)
        prio_cell.fill      = bg_fill
        prio_cell.border    = thin_border()
        prio_cell.alignment = center_align()

        verif_cell = ws2.cell(row=row_num, column=4, value=xl_ef_verif(exig))
        verif_cell.font      = body_font(size=9, color=CYAN_HEX)
        verif_cell.fill      = bg_fill
        verif_cell.border    = thin_border()
        verif_cell.alignment = center_align()

        stat_cell = ws2.cell(row=row_num, column=5, value="À valider")
        stat_cell.font      = body_font(size=9, color="666666")
        stat_cell.fill      = bg_fill
        stat_cell.border    = thin_border()
        stat_cell.alignment = center_align()

        ws2.row_dimensions[row_num].height = 32

    # Dropdowns priorité, vérification, statut
    last_ef = len(ef) + 2
    dv_prio = DataValidation(type="list", formula1='"Must,Should,Could,Won\'t"',
                             allow_blank=True, showDropDown=False)
    dv_verif = DataValidation(type="list",
                              formula1='"Test,Analyse,Inspection,Démonstration"',
                              allow_blank=True, showDropDown=False)
    dv_stat = DataValidation(type="list",
                             formula1='"À valider,Validé,Rejeté,En révision"',
                             allow_blank=True, showDropDown=False)
    ws2.add_data_validation(dv_prio);  dv_prio.sqref  = f"C3:C{last_ef}"
    ws2.add_data_validation(dv_verif); dv_verif.sqref = f"D3:D{last_ef}"
    ws2.add_data_validation(dv_stat);  dv_stat.sqref  = f"E3:E{last_ef}"

    ws2.freeze_panes = "A3"

    # =========================================================
    # FEUILLE 3 : CONTRAINTES TECHNIQUES
    # =========================================================
    ws3 = wb.create_sheet("Contraintes Techniques")
    ws3.sheet_view.showGridLines = False

    ws3.merge_cells("A1:B1")
    ct_title = ws3["A1"]
    ct_title.value     = f"CONTRAINTES TECHNIQUES — {p.get('systeme_sous_analyse','')}"
    ct_title.font      = Font(name="Arial", bold=True, size=13, color=NAVY_HEX)
    ct_title.fill      = light_fill(LIGHT_BG)
    ct_title.alignment = center_align()
    ws3.row_dimensions[1].height = 26

    apply_header_row(ws3, 2,
        ["DOMAINE DE CONTRAINTE", "SPÉCIFICATION"],
        col_widths=[28, 80])

    CT_LABELS = [
        ("acoustique",             "Acoustique"),
        ("frequences",             "Fréquences (VLF/ELF/…)"),
        ("etancheite_pression",    "Étanchéité / Pression"),
        ("debit_latence",          "Débit / Latence"),
        ("environnement_physique", "Environnement physique"),
    ]
    for idx, (key, label) in enumerate(CT_LABELS):
        row_num = idx + 3
        val     = ct.get(key, "Non spécifié")
        is_na   = val in ("N/A","Non spécifié","Non specifie")
        bg_fill = light_fill(LIGHT_BG) if idx % 2 == 0 else alt_fill()

        lbl_c = ws3.cell(row=row_num, column=1, value=label)
        lbl_c.font      = body_font(bold=True, color=NAVY_HEX)
        lbl_c.fill      = bg_fill
        lbl_c.border    = thin_border()
        lbl_c.alignment = left_align()

        val_c = ws3.cell(row=row_num, column=2, value=val)
        val_c.font      = body_font(size=10, color="AAAAAA" if is_na else "222222",
                                    italic=is_na)
        val_c.fill      = bg_fill
        val_c.border    = thin_border()
        val_c.alignment = left_align(wrap=True)
        ws3.row_dimensions[row_num].height = 28

    # Autres contraintes
    autres = ct.get("autres", [])
    if autres:
        row_num = len(CT_LABELS) + 3
        bg_fill = light_fill(LIGHT_BG) if len(CT_LABELS) % 2 == 0 else alt_fill()
        lbl_c = ws3.cell(row=row_num, column=1, value="Autres contraintes")
        lbl_c.font = body_font(bold=True, color=NAVY_HEX)
        lbl_c.fill = bg_fill; lbl_c.border = thin_border()
        lbl_c.alignment = left_align()
        val_c = ws3.cell(row=row_num, column=2, value=" / ".join(autres))
        val_c.font = body_font(size=10)
        val_c.fill = bg_fill; val_c.border = thin_border()
        val_c.alignment = left_align(wrap=True)
        ws3.row_dimensions[row_num].height = 28

    ws3.freeze_panes = "A3"

    # =========================================================
    # FEUILLE 4 : EXIGENCES NON-FONCTIONNELLES
    # =========================================================
    nfr = result.get("exigences_non_fonctionnelles", [])
    if nfr:
        ws4 = wb.create_sheet("Exigences Non-Fonctionnelles")
        ws4.sheet_view.showGridLines = False

        ws4.merge_cells("A1:E1")
        nfr_title = ws4["A1"]
        nfr_title.value     = f"EXIGENCES NON-FONCTIONNELLES — {p.get('systeme_sous_analyse','')}"
        nfr_title.font      = Font(name="Arial", bold=True, size=13, color=NAVY_HEX)
        nfr_title.fill      = light_fill(LIGHT_BG)
        nfr_title.alignment = center_align()
        ws4.row_dimensions[1].height = 26

        apply_header_row(ws4, 2,
            ["ID", "EXIGENCE", "TYPE", "PRIORITÉ", "VÉRIF."],
            col_widths=[8, 52, 14, 12, 16])

        for idx, item in enumerate(nfr):
            row_num  = idx + 3
            bg_fill  = light_fill(LIGHT_BG) if idx % 2 == 0 else alt_fill()
            nfr_text  = item.get("text", item.get("texte", "")) if isinstance(item, dict) else str(item)
            nfr_type  = item.get("type", "Performance") if isinstance(item, dict) else "Performance"
            nfr_prio  = item.get("priority", "Must") if isinstance(item, dict) else "Must"
            nfr_verif = item.get("verification", "Test") if isinstance(item, dict) else "Test"

            c1 = ws4.cell(row=row_num, column=1, value=f"ENF-{str(idx+1).zfill(2)}")
            c1.font = body_font(bold=True, color=CYAN_HEX); c1.fill = bg_fill
            c1.border = thin_border(); c1.alignment = center_align()

            c2 = ws4.cell(row=row_num, column=2, value=nfr_text)
            c2.font = body_font(size=10); c2.fill = bg_fill
            c2.border = thin_border(); c2.alignment = left_align(wrap=True)

            c3 = ws4.cell(row=row_num, column=3, value=nfr_type)
            c3.font = body_font(size=9, color=ORANGE_HEX, bold=True); c3.fill = bg_fill
            c3.border = thin_border(); c3.alignment = center_align()

            c4 = ws4.cell(row=row_num, column=4, value=nfr_prio)
            c4.font = body_font(size=9, color=PRIO_COLORS_XL.get(nfr_prio, ORANGE_HEX), bold=True)
            c4.fill = bg_fill; c4.border = thin_border(); c4.alignment = center_align()

            c5 = ws4.cell(row=row_num, column=5, value=nfr_verif)
            c5.font = body_font(size=9, color=CYAN_HEX); c5.fill = bg_fill
            c5.border = thin_border(); c5.alignment = center_align()

            ws4.row_dimensions[row_num].height = 32

        ws4.freeze_panes = "A3"

        # Dropdowns NFR
        last_nfr = len(nfr) + 2
        dv_nfr_prio = DataValidation(type="list", formula1='"Must,Should,Could,Won\'t"',
                                     allow_blank=True, showDropDown=False)
        dv_nfr_verif = DataValidation(type="list",
                                      formula1='"Test,Analyse,Inspection,Démonstration"',
                                      allow_blank=True, showDropDown=False)
        ws4.add_data_validation(dv_nfr_prio);  dv_nfr_prio.sqref  = f"D3:D{last_nfr}"
        ws4.add_data_validation(dv_nfr_verif); dv_nfr_verif.sqref = f"E3:E{last_nfr}"

    # =========================================================
    # FEUILLE 5 : MATRICE DE TRAÇABILITÉ
    # =========================================================
    ws5 = wb.create_sheet("Matrice Traçabilité")
    ws5.sheet_view.showGridLines = False

    ws5.merge_cells("A1:H1")
    tr_title = ws5["A1"]
    tr_title.value     = f"MATRICE DE TRAÇABILITÉ — {p.get('systeme_sous_analyse','')}"
    tr_title.font      = Font(name="Arial", bold=True, size=13, color=NAVY_HEX)
    tr_title.fill      = light_fill(LIGHT_BG)
    tr_title.alignment = center_align()
    ws5.row_dimensions[1].height = 26

    # Texte source en ligne 2
    demande_src = meta.get("demande_full") or meta.get("demande_source") or ""
    if demande_src:
        ws5.merge_cells("A2:H2")
        src_cell = ws5["A2"]
        src_cell.value     = f"Source : {demande_src[:300]}{'...' if len(demande_src) > 300 else ''}"
        src_cell.font      = Font(name="Arial", size=8, color="889999", italic=True)
        src_cell.fill      = PatternFill("solid", fgColor=LIGHT_BG)
        src_cell.alignment = left_align(wrap=True)
        ws5.row_dimensions[2].height = 30

    header_row = 3 if demande_src else 2
    apply_header_row(ws5, header_row,
        ["ID", "EXIGENCE", "TYPE", "PRIORITÉ", "MÉTH. VÉRIF.", "RÉFÉRENCE SOURCE", "RESPONSABLE", "STATUT"],
        col_widths=[8, 50, 14, 12, 16, 20, 16, 14])

    all_items = (
        [("EF",  str(i+1).zfill(2), xl_ef_text(e),  "Fonctionnelle",
          xl_ef_prio(e), xl_ef_verif(e)) for i, e in enumerate(ef)] +
        [("ENF", str(i+1).zfill(2),
          item.get("text", item.get("texte","")) if isinstance(item, dict) else str(item),
          item.get("type","Non-Fonctionnelle") if isinstance(item, dict) else "Non-Fonctionnelle",
          item.get("priority","Must") if isinstance(item, dict) else "Must",
          item.get("verification","Test") if isinstance(item, dict) else "Test")
         for i, item in enumerate(nfr if nfr else [])]
    )

    dv_tr_prio  = DataValidation(type="list", formula1='"Must,Should,Could,Won\'t"',
                                 allow_blank=True, showDropDown=False)
    dv_tr_verif = DataValidation(type="list",
                                 formula1='"Test,Analyse,Inspection,Démonstration"',
                                 allow_blank=True, showDropDown=False)
    dv_tr_stat  = DataValidation(type="list",
                                 formula1='"À valider,Validé,Rejeté,En révision"',
                                 allow_blank=True, showDropDown=False)
    ws5.add_data_validation(dv_tr_prio)
    ws5.add_data_validation(dv_tr_verif)
    ws5.add_data_validation(dv_tr_stat)

    for idx, (prefix, num, text_val, type_val, prio_val, verif_val) in enumerate(all_items):
        rn = idx + header_row + 1
        bg = light_fill(LIGHT_BG) if idx % 2 == 0 else alt_fill()

        c1 = ws5.cell(row=rn, column=1, value=f"{prefix}-{num}")
        c1.font = body_font(bold=True, color=CYAN_HEX); c1.fill = bg
        c1.border = thin_border(); c1.alignment = center_align()

        c2 = ws5.cell(row=rn, column=2, value=text_val)
        c2.font = body_font(size=10); c2.fill = bg
        c2.border = thin_border(); c2.alignment = left_align(wrap=True)

        c3 = ws5.cell(row=rn, column=3, value=type_val)
        c3.font = body_font(size=9, color=ORANGE_HEX, bold=True); c3.fill = bg
        c3.border = thin_border(); c3.alignment = center_align()

        c4 = ws5.cell(row=rn, column=4, value=prio_val)
        c4.font = body_font(size=9, color=PRIO_COLORS_XL.get(prio_val, ORANGE_HEX), bold=True)
        c4.fill = bg; c4.border = thin_border(); c4.alignment = center_align()

        c5 = ws5.cell(row=rn, column=5, value=verif_val)
        c5.font = body_font(size=9, color=CYAN_HEX); c5.fill = bg
        c5.border = thin_border(); c5.alignment = center_align()

        for col in (6, 7):  # Référence source + Responsable (à remplir)
            cx = ws5.cell(row=rn, column=col, value="")
            cx.fill = bg; cx.border = thin_border(); cx.alignment = left_align()

        c8 = ws5.cell(row=rn, column=8, value="À valider")
        c8.font = body_font(size=9, color="666666"); c8.fill = bg
        c8.border = thin_border(); c8.alignment = center_align()
        ws5.row_dimensions[rn].height = 30

    last_tr = len(all_items) + header_row
    dv_tr_prio.sqref  = f"D{header_row+1}:D{last_tr}"
    dv_tr_verif.sqref = f"E{header_row+1}:E{last_tr}"
    dv_tr_stat.sqref  = f"H{header_row+1}:H{last_tr}"
    ws5.freeze_panes  = f"A{header_row+1}"

    # =========================================================
    # Sérialisation
    # =========================================================
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"exigences_{ts}.xlsx"

    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ===========================================================================
# EXPORT DOORS CSV (IBM DOORS Next / DNG)
# ===========================================================================
@app.route("/export/doors-csv", methods=["POST"])
def export_doors_csv():
    """Export CSV compatible IBM DOORS Next (import via assistant de mappage)."""
    import csv
    data   = request.get_json(force=True)
    result = data.get("result")
    if not result:
        return jsonify({"error": "Rien à exporter."}), 400

    p       = result.get("perimetre", {})
    ef      = result.get("exigences_fonctionnelles", [])
    nfr     = result.get("exigences_non_fonctionnelles", [])
    meta    = result.get("_meta", {})
    systeme = p.get("systeme_sous_analyse", "Système")
    besoin  = result.get("besoin_fondamental", "")

    def ef_t(e): return e.get("text", e) if isinstance(e, dict) else e
    def ef_p(e): return e.get("priority", "Must") if isinstance(e, dict) else "Must"
    def ef_v(e): return e.get("verification", "Test") if isinstance(e, dict) else "Test"

    output  = io.StringIO()
    writer  = csv.writer(output, quoting=csv.QUOTE_ALL)

    # En-têtes DOORS Next compatibles
    writer.writerow([
        "Identifier", "Artifact Type", "Module", "Primary Text",
        "Priority", "Verification Method", "Status",
        "System", "Fundamental Need", "Created By", "Created On"
    ])

    ts_str = ""
    if meta.get("timestamp"):
        try: ts_str = datetime.datetime.fromisoformat(meta["timestamp"]).strftime("%d/%m/%Y %H:%M")
        except Exception: ts_str = meta.get("timestamp", "")

    # EF
    for idx, exig in enumerate(ef):
        writer.writerow([
            f"EF-{str(idx+1).zfill(2)}",
            "Requirement",
            "Exigences Fonctionnelles",
            ef_t(exig),
            ef_p(exig),
            ef_v(exig),
            "Proposed",
            systeme,
            besoin,
            "SYREQ",
            ts_str
        ])

    # ENF
    for idx, item in enumerate(nfr):
        nfr_text  = item.get("text", item.get("texte", "")) if isinstance(item, dict) else str(item)
        nfr_type  = item.get("type", "Performance") if isinstance(item, dict) else "Performance"
        nfr_prio  = item.get("priority", "Must") if isinstance(item, dict) else "Must"
        nfr_verif = item.get("verification", "Test") if isinstance(item, dict) else "Test"
        writer.writerow([
            f"ENF-{str(idx+1).zfill(2)}",
            "Requirement",
            f"Exigences Non-Fonctionnelles — {nfr_type}",
            nfr_text,
            nfr_prio,
            nfr_verif,
            "Proposed",
            systeme,
            besoin,
            "SYREQ",
            ts_str
        ])

    csv_bytes = output.getvalue().encode("utf-8-sig")  # BOM pour Excel/DOORS
    buf = io.BytesIO(csv_bytes)
    buf.seek(0)

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    return send_file(buf, as_attachment=True,
                     download_name=f"exigences_{ts}_doors.csv",
                     mimetype="text/csv")


# ===========================================================================
# EXPORT ReqIF (IBM DOORS Next / Jazz ELM — standard OMG)
# ===========================================================================
@app.route("/export/reqif", methods=["POST"])
def export_reqif():
    """Export ReqIF 1.0.1 — format natif IBM DOORS Next et Jazz ELM."""
    import xml.sax.saxutils as saxutils

    data   = request.get_json(force=True)
    result = data.get("result")
    if not result:
        return jsonify({"error": "Rien à exporter."}), 400

    p       = result.get("perimetre", {})
    ef      = result.get("exigences_fonctionnelles", [])
    nfr     = result.get("exigences_non_fonctionnelles", [])
    meta    = result.get("_meta", {})
    systeme = p.get("systeme_sous_analyse", "Système")
    now_iso = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%S")

    def esc(s):   return saxutils.escape(str(s))
    def slug(s):  return re.sub(r"[^A-Za-z0-9_-]", "-", str(s))[:40]
    def ef_t(e):  return e.get("text", e) if isinstance(e, dict) else e
    def ef_p(e):  return e.get("priority", "Must") if isinstance(e, dict) else "Must"
    def ef_v(e):  return e.get("verification", "Test") if isinstance(e, dict) else "Test"

    prio_map  = {"Must": "EV-MUST", "Should": "EV-SHOULD",
                 "Could": "EV-COULD", "Won't": "EV-WONT"}
    verif_map = {"Test": "EV-TEST", "Analyse": "EV-ANALYSE",
                 "Inspection": "EV-INSP", "Démonstration": "EV-DEMO"}

    all_reqs = []
    for i, e in enumerate(ef):
        all_reqs.append({
            "id": f"SO-EF-{str(i+1).zfill(2)}",
            "name": f"EF-{str(i+1).zfill(2)}",
            "text": ef_t(e),
            "type": "Fonctionnelle",
            "prio": ef_p(e),
            "verif": ef_v(e),
            "module": "Exigences Fonctionnelles"
        })
    for i, item in enumerate(nfr):
        nfr_t = item.get("text", item.get("texte","")) if isinstance(item,dict) else str(item)
        nfr_type = item.get("type","Performance") if isinstance(item,dict) else "Performance"
        all_reqs.append({
            "id": f"SO-ENF-{str(i+1).zfill(2)}",
            "name": f"ENF-{str(i+1).zfill(2)}",
            "text": nfr_t,
            "type": nfr_type,
            "prio": item.get("priority","Must") if isinstance(item,dict) else "Must",
            "verif": item.get("verification","Test") if isinstance(item,dict) else "Test",
            "module": f"Exigences Non-Fonctionnelles — {nfr_type}"
        })

    lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<REQ-IF xmlns="http://www.omg.org/spec/ReqIF/20110401/reqif.xsd"',
        '        xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">',
        '  <THE-HEADER>',
        '    <REQ-IF-HEADER IDENTIFIER="RIF-HDR-SYREQ">',
        f'      <CREATION-TIME>{now_iso}</CREATION-TIME>',
        '      <REQ-IF-TOOL-ID>SYREQ v1.0</REQ-IF-TOOL-ID>',
        '      <REQ-IF-VERSION>1.0.1</REQ-IF-VERSION>',
        '      <SOURCE-TOOL-ID>SYREQ — Analyseur d\'Exigences Naval</SOURCE-TOOL-ID>',
        f'      <TITLE>{esc(systeme)}</TITLE>',
        '    </REQ-IF-HEADER>',
        '  </THE-HEADER>',
        '  <CORE-CONTENT><REQ-IF-CONTENT>',
        '    <DATATYPES>',
        f'      <DATATYPE-DEFINITION-STRING IDENTIFIER="DT-STRING" LAST-CHANGE="{now_iso}" LONG-NAME="String" MAX-LENGTH="50000"/>',
        f'      <DATATYPE-DEFINITION-ENUMERATION IDENTIFIER="DT-PRIORITY" LAST-CHANGE="{now_iso}" LONG-NAME="Priority">',
        '        <SPECIFIED-VALUES>',
        f'          <ENUM-VALUE IDENTIFIER="EV-MUST"   LAST-CHANGE="{now_iso}" LONG-NAME="Must"/>',
        f'          <ENUM-VALUE IDENTIFIER="EV-SHOULD" LAST-CHANGE="{now_iso}" LONG-NAME="Should"/>',
        f'          <ENUM-VALUE IDENTIFIER="EV-COULD"  LAST-CHANGE="{now_iso}" LONG-NAME="Could"/>',
        f'          <ENUM-VALUE IDENTIFIER="EV-WONT"   LAST-CHANGE="{now_iso}" LONG-NAME="Won\'t"/>',
        '        </SPECIFIED-VALUES>',
        '      </DATATYPE-DEFINITION-ENUMERATION>',
        f'      <DATATYPE-DEFINITION-ENUMERATION IDENTIFIER="DT-VERIF" LAST-CHANGE="{now_iso}" LONG-NAME="Verification Method">',
        '        <SPECIFIED-VALUES>',
        f'          <ENUM-VALUE IDENTIFIER="EV-TEST"    LAST-CHANGE="{now_iso}" LONG-NAME="Test"/>',
        f'          <ENUM-VALUE IDENTIFIER="EV-ANALYSE" LAST-CHANGE="{now_iso}" LONG-NAME="Analyse"/>',
        f'          <ENUM-VALUE IDENTIFIER="EV-INSP"    LAST-CHANGE="{now_iso}" LONG-NAME="Inspection"/>',
        f'          <ENUM-VALUE IDENTIFIER="EV-DEMO"    LAST-CHANGE="{now_iso}" LONG-NAME="Démonstration"/>',
        '        </SPECIFIED-VALUES>',
        '      </DATATYPE-DEFINITION-ENUMERATION>',
        '    </DATATYPES>',
        '    <SPEC-TYPES>',
        f'      <SPEC-OBJECT-TYPE IDENTIFIER="SOT-REQ" LAST-CHANGE="{now_iso}" LONG-NAME="Exigence">',
        '        <SPEC-ATTRIBUTES>',
        f'          <ATTRIBUTE-DEFINITION-STRING IDENTIFIER="AD-TEXT"   LAST-CHANGE="{now_iso}" LONG-NAME="Texte"><TYPE><DATATYPE-DEFINITION-STRING-REF>DT-STRING</DATATYPE-DEFINITION-STRING-REF></TYPE></ATTRIBUTE-DEFINITION-STRING>',
        f'          <ATTRIBUTE-DEFINITION-STRING IDENTIFIER="AD-TYPE"   LAST-CHANGE="{now_iso}" LONG-NAME="Type"><TYPE><DATATYPE-DEFINITION-STRING-REF>DT-STRING</DATATYPE-DEFINITION-STRING-REF></TYPE></ATTRIBUTE-DEFINITION-STRING>',
        f'          <ATTRIBUTE-DEFINITION-STRING IDENTIFIER="AD-MODULE" LAST-CHANGE="{now_iso}" LONG-NAME="Module"><TYPE><DATATYPE-DEFINITION-STRING-REF>DT-STRING</DATATYPE-DEFINITION-STRING-REF></TYPE></ATTRIBUTE-DEFINITION-STRING>',
        f'          <ATTRIBUTE-DEFINITION-ENUMERATION IDENTIFIER="AD-PRIO"  LAST-CHANGE="{now_iso}" LONG-NAME="Priorité" MULTI-VALUED="false"><TYPE><DATATYPE-DEFINITION-ENUMERATION-REF>DT-PRIORITY</DATATYPE-DEFINITION-ENUMERATION-REF></TYPE></ATTRIBUTE-DEFINITION-ENUMERATION>',
        f'          <ATTRIBUTE-DEFINITION-ENUMERATION IDENTIFIER="AD-VERIF" LAST-CHANGE="{now_iso}" LONG-NAME="Vérification" MULTI-VALUED="false"><TYPE><DATATYPE-DEFINITION-ENUMERATION-REF>DT-VERIF</DATATYPE-DEFINITION-ENUMERATION-REF></TYPE></ATTRIBUTE-DEFINITION-ENUMERATION>',
        '        </SPEC-ATTRIBUTES>',
        '      </SPEC-OBJECT-TYPE>',
        f'      <SPECIFICATION-TYPE IDENTIFIER="ST-SPEC" LAST-CHANGE="{now_iso}" LONG-NAME="Spécification"/>',
        '    </SPEC-TYPES>',
        '    <SPEC-OBJECTS>',
    ]

    for req in all_reqs:
        prio_ref  = prio_map.get(req["prio"],  "EV-MUST")
        verif_ref = verif_map.get(req["verif"], "EV-TEST")
        lines += [
            f'      <SPEC-OBJECT IDENTIFIER="{req["id"]}" LAST-CHANGE="{now_iso}" LONG-NAME="{esc(req["name"])}">',
            '        <TYPE><SPEC-OBJECT-TYPE-REF>SOT-REQ</SPEC-OBJECT-TYPE-REF></TYPE>',
            '        <VALUES>',
            f'          <ATTRIBUTE-VALUE-STRING THE-VALUE="{esc(req["text"])}"><DEFINITION><ATTRIBUTE-DEFINITION-STRING-REF>AD-TEXT</ATTRIBUTE-DEFINITION-STRING-REF></DEFINITION></ATTRIBUTE-VALUE-STRING>',
            f'          <ATTRIBUTE-VALUE-STRING THE-VALUE="{esc(req["type"])}"><DEFINITION><ATTRIBUTE-DEFINITION-STRING-REF>AD-TYPE</ATTRIBUTE-DEFINITION-STRING-REF></DEFINITION></ATTRIBUTE-VALUE-STRING>',
            f'          <ATTRIBUTE-VALUE-STRING THE-VALUE="{esc(req["module"])}"><DEFINITION><ATTRIBUTE-DEFINITION-STRING-REF>AD-MODULE</ATTRIBUTE-DEFINITION-STRING-REF></DEFINITION></ATTRIBUTE-VALUE-STRING>',
            f'          <ATTRIBUTE-VALUE-ENUMERATION><DEFINITION><ATTRIBUTE-DEFINITION-ENUMERATION-REF>AD-PRIO</ATTRIBUTE-DEFINITION-ENUMERATION-REF></DEFINITION><VALUES><ENUM-VALUE-REF>{prio_ref}</ENUM-VALUE-REF></VALUES></ATTRIBUTE-VALUE-ENUMERATION>',
            f'          <ATTRIBUTE-VALUE-ENUMERATION><DEFINITION><ATTRIBUTE-DEFINITION-ENUMERATION-REF>AD-VERIF</ATTRIBUTE-DEFINITION-ENUMERATION-REF></DEFINITION><VALUES><ENUM-VALUE-REF>{verif_ref}</ENUM-VALUE-REF></VALUES></ATTRIBUTE-VALUE-ENUMERATION>',
            '        </VALUES>',
            '      </SPEC-OBJECT>',
        ]

    lines += [
        '    </SPEC-OBJECTS>',
        '    <SPEC-RELATIONS/>',
        '    <SPECIFICATIONS>',
        f'      <SPECIFICATION IDENTIFIER="SPEC-1" LAST-CHANGE="{now_iso}" LONG-NAME="{esc(systeme)}">',
        '        <TYPE><SPECIFICATION-TYPE-REF>ST-SPEC</SPECIFICATION-TYPE-REF></TYPE>',
        '        <CHILDREN>',
    ]
    for req in all_reqs:
        lines.append(f'          <SPEC-HIERARCHY IDENTIFIER="SH-{slug(req["id"])}" LAST-CHANGE="{now_iso}"><OBJECT><SPEC-OBJECT-REF>{req["id"]}</SPEC-OBJECT-REF></OBJECT></SPEC-HIERARCHY>')
    lines += [
        '        </CHILDREN>',
        '      </SPECIFICATION>',
        '    </SPECIFICATIONS>',
        '  </REQ-IF-CONTENT></CORE-CONTENT>',
        '</REQ-IF>',
    ]

    xml_bytes = "\n".join(lines).encode("utf-8")
    buf = io.BytesIO(xml_bytes)
    buf.seek(0)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    return send_file(buf, as_attachment=True,
                     download_name=f"exigences_{ts}.reqif",
                     mimetype="application/xml")


# ---------------------------------------------------------------------------
# /models
# ---------------------------------------------------------------------------
@app.route("/models", methods=["GET"])
def list_models():
    names = []

    # Modèles Ollama locaux
    try:
        models_response = ollama.list()
        raw_models = (
            models_response.models
            if hasattr(models_response, "models")
            else models_response.get("models", [])
        )
        for m in raw_models:
            name = (
                getattr(m, "model", None)
                or getattr(m, "name", None)
                or (m.get("name") if isinstance(m, dict) else None)
                or (m.get("model") if isinstance(m, dict) else None)
                or ""
            )
            if name:
                names.append(name)
    except Exception:
        if not names:
            names += ["llama3", "mistral"]

    return jsonify({"models": names if names else ["llama3", "mistral"]})


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("=" * 62)
    print("   Analyseur d'Exigences — Domaine Naval / Sous-marin")
    print("=" * 62)
    print(f"   Modèle par défaut  : {DEFAULT_MODEL}")
    print(f"   Dossier exports    : ./{OUTPUT_DIR}/")
    print(f"   Interface web      : http://{HOST}:{PORT}")
    print("=" * 62)
    try:
        ollama.list()
        print("   [OK] Ollama est accessible.\n")
    except Exception:
        print("   [ATTENTION] Ollama ne semble pas démarré.\n")
    app.run(host=HOST, port=PORT, debug=False)
